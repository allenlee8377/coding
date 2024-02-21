from docx import Document
import openpyxl

def convert_word_to_excel(word_path, excel_path):
    # 打开 Word 文档
    doc = Document(word_path)

    # 创建 Excel 文件
    wb = openpyxl.Workbook()
    ws = wb.active

    for paragraph in doc.paragraphs:
        # 将段落文本写入 Excel 文件
        ws.append([paragraph.text])

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            ws.append(row_data)

    # 保存 Excel 文件
    wb.save(excel_path)
    print(f"成功将Word文档转换为Excel文件: {excel_path}")

# 输入 Word 和 Excel 文件的路径
word_file_path = input("请输入Word文档路径：")
excel_file_path = input("请输入要保存的Excel文件路径：")

# 调用转换函数
convert_word_to_excel(word_file_path, excel_file_path)
